import io
import json
import os
import sqlite3
import time
import bcrypt
import docx
from fpdf import FPDF
import google.generativeai as genai
from PyPDF2 import PdfReader
import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
from typing import List, Dict
from mistralai import Mistral
import requests
import pytesseract
from pdf2image import convert_from_bytes
import re
from openai import OpenAI
from PIL import Image
import pdfplumber
import torch
from torchvision import transforms
import uuid

# Configuration
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY", "AIzaSyANbVVzZACnYnus00xwwRRE01n34yoAmcU")
MISTRAL_API_KEY = os.environ.get("MISTRAL_API_KEY", "DUW9f3t6nvZaNkEbxcrxYP4hLIrC3g7Y")
MISTRAL_ENDPOINT = "https://api.mistral.ai/v1/chat/completions"
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "sk-61f7f17d33bd4598b4dd61edd13af337")
OCR_API_KEY = "K85749057588957"  # Your OCR.space API key

DEEPSEEK_CLIENT = OpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com/v1"
)

BRAND_COLORS = {
    "primary": "#2E86AB",
    "secondary": "#F18F01",
    "background": "#F7F7F7",
    "text": "#121111"
}

genai.configure(api_key=GOOGLE_API_KEY)
mistral_client = Mistral(api_key=MISTRAL_API_KEY)

# Database Helpers
def init_db(db_path: str = "users.db"):
    conn = sqlite3.connect(db_path, check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password BLOB NOT NULL,
            role TEXT NOT NULL,
            location_id TEXT,
            last_login TIMESTAMP,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    cursor.execute("PRAGMA table_info(users)")
    existing_cols = [col[1] for col in cursor.fetchall()]
    for col in ["location_id", "last_login", "created_at"]:
        if col not in existing_cols:
            try:
                cursor.execute(f"ALTER TABLE users ADD COLUMN {col} TEXT")
            except sqlite3.OperationalError:
                pass

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS subscriptions (
            username TEXT PRIMARY KEY,
            lease_analysis BOOLEAN DEFAULT 0,
            deal_structuring BOOLEAN DEFAULT 0,
            offer_generator BOOLEAN DEFAULT 0,
            FOREIGN KEY(username) REFERENCES users(username)
        )
    """)

    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS interactions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT,
            feature TEXT,
            input_text TEXT,
            output_text TEXT,
            timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY(username) REFERENCES users(username)
        )
        """
    )
    conn.commit()
    return conn

def create_default_admin(conn):
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) FROM users")
    if cursor.fetchone()[0] == 0:
        admin_pwd = bcrypt.hashpw("admin123".encode(), bcrypt.gensalt())
        cursor.execute(
            "INSERT INTO users (username, password, role) VALUES (?, ?, ?)",
            ("admin", admin_pwd, "admin"),
        )
        conn.commit()

def verify_password(hashed: bytes, password: str) -> bool:
    return bcrypt.checkpw(password.encode(), hashed)

# Authentication UI
def login_ui(conn):
    st.sidebar.title("🔑 Login / Register")
    st.sidebar.markdown(
        """
        <style>
            .sidebar .sidebar-content {
                background-color: #FFFFFF !important;
                color: #000000 !important;
                box-shadow: none !important;
            }
            .stTextInput>div>div>input,
            .stTextArea>div>div>textarea {
                background-color: #FFFFFF !important;
                color: #000000 !important;
                border: 1px solid #CCCCCC !important;
            }
            .stButton>button {
                background-color: #FFFFFF !important;
                color: #000000 !important;
                border: 1px solid #CCCCCC !important;
            }
        </style>
        """,
        unsafe_allow_html=True
    )

    login_tab, register_tab = st.sidebar.tabs(["Login", "Register"])

    with login_tab:
        username = st.text_input("Username", key="login_username")
        password = st.text_input("Password", type="password", key="login_password")
        location_id = st.text_input("Login Key (use instead of user/pass)", key="login_location")

        if st.button("Log In", key="login_button"):
            if location_id and not username and not password:
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT username, role FROM users WHERE location_id = ?",
                    (location_id,)
                )
                row = cursor.fetchone()
                if row:
                    st.session_state.logged_in = True
                    st.session_state.username = row[0]
                    st.session_state.role = row[1]
                    st.session_state.location_id = location_id
                    cursor.execute("UPDATE users SET last_login = CURRENT_TIMESTAMP WHERE username = ?", (row[0],))
                    conn.commit()
                    st.rerun()
                else:
                    st.sidebar.error("❌ Invalid login key.")
            elif not username or not password:
                st.sidebar.error("Enter both username and password")
            else:
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT password, role, location_id FROM users WHERE username = ?",
                    (username,)
                )
                row = cursor.fetchone()
                if row and bcrypt.checkpw(password.encode(), row[0]):
                    st.session_state.logged_in = True
                    st.session_state.username = username
                    st.session_state.role = row[1]
                    st.session_state.location_id = row[2]
                    cursor.execute("UPDATE users SET last_login = CURRENT_TIMESTAMP WHERE username = ?", (username,))
                    conn.commit()
                    st.rerun()
                else:
                    st.sidebar.error("Invalid username or password")
                    time.sleep(1)

    with register_tab:
        new_user = st.text_input("New Username", key="reg_username")
        new_pass = st.text_input("New Password", type="password", key="reg_password")
        confirm_pass = st.text_input("Confirm Password", type="password", key="reg_confirm")
        user_role = "user"

        if st.button("Create User", key="reg_button"):
            if not new_user or not new_pass:
                st.error("Username and password are required")
            elif new_pass != confirm_pass:
                st.error("Passwords do not match")
            elif len(new_pass) < 8:
                st.error("Password must be at least 8 characters")
            else:
                try:
                    location_key = str(uuid.uuid4())
                    hashed = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt())
                    cursor = conn.cursor()
                    cursor.execute(
                        "INSERT INTO users (username, password, role, location_id) VALUES (?, ?, ?, ?)",
                        (new_user, hashed, user_role, location_key)
                    )
                    conn.commit()
                    st.success(f"User '{new_user}' created successfully.")
                    st.info(f"🔑 **Your login key** (save this!): `{location_key}`")
                    time.sleep(1)
                    st.rerun()
                except sqlite3.IntegrityError:
                    st.error("That username already exists")

    if not st.session_state.get("logged_in"):
        st.markdown(
            """
            <div style="
                height: 80vh;
                display: flex;
                flex-direction: column;
                align-items: center;
                justify-content: center;
                text-align: center;
            ">
                <h1 style="color: #2E86AB; font-weight: normal; margin-bottom: 0.2em;">
                    Welcome to Finchat AI Bot
                </h1>
                <p style="color: #555555; font-size: 1.1em; margin-top: 0;">
                    🤖 Powered by Alphax — crafting real estate insights in seconds!
                </p>
            </div>
            """,
            unsafe_allow_html=True
        )
        return

# OCR Functions
def ocr_space_file(filename, file_bytes, overlay=False, language='eng'):
    """OCR.space API request with file bytes."""
    payload = {
        'isOverlayRequired': overlay,
        'apikey': OCR_API_KEY,
        'language': language,
    }

    files = {filename: file_bytes}
    response = requests.post(
        'https://api.ocr.space/parse/image',
        files=files,
        data=payload,
    )

    result = response.json()
    if response.status_code == 200 and result.get('IsErroredOnProcessing', True) == False:
        return result['ParsedResults'][0]['ParsedText']
    else:
        error_message = result.get('ErrorMessage', 'Unknown error occurred')
        raise Exception(f"OCR processing failed: {error_message}")

def is_scanned_pdf(pdf_bytes):
    """Check if PDF is scanned (non-selectable text)."""
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and len(text.strip()) > 50:  # If we find reasonable text
                    return False
        return True
    except:
        return True

def process_pdf_with_ocr(pdf_bytes, filename):
    """Process PDF with OCR if needed."""
    if is_scanned_pdf(pdf_bytes):
        st.info("🔍 Detected scanned document - performing OCR...")

        # Convert PDF to images and OCR each page
        images = convert_from_bytes(pdf_bytes)
        all_text = []

        for i, image in enumerate(images):
            with st.spinner(f"Processing page {i+1}/{len(images)}..."):
                # Save image to bytes
                img_byte_arr = io.BytesIO()
                image.save(img_byte_arr, format='PNG')
                img_byte_arr = img_byte_arr.getvalue()

                # Call OCR API
                try:
                    text = ocr_space_file(f"page_{i+1}.png", img_byte_arr)
                    all_text.append(text)
                except Exception as e:
                    st.error(f"Error processing page {i+1}: {str(e)}")
                    all_text.append("(OCR failed for this page)")

        return "\n\n".join(all_text)
    else:
        # Extract text directly if not scanned
        reader = PdfReader(io.BytesIO(pdf_bytes))
        return "\n".join([page.extract_text() or "" for page in reader.pages])

# AI Helper Functions
def call_gemini(feature: str, content: str, temperature: float = 0.7) -> str:
    system_prompts = {
        "lease_analysis": (
            "You are a real estate document expert. Analyze the provided lease agreement "
            "and provide a comprehensive summary, including key terms and potential risks."
        ),
        "deal_strategy": (
            "You are a creative real estate strategist. Based on the provided deal details, "
            "suggest structuring options with pros, cons, and negotiation tactics."
        ),
        "offer_generator": (
            "You are a real estate transaction specialist. Generate a professional purchase offer "
            "with all essential clauses formatted for the jurisdiction."
        ),
        "chatbot": (
            "You are a knowledgeable assistant that answers questions based on the user's past interactions."
        )
    }

    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"SYSTEM: {system_prompts.get(feature, '')}\n\nUSER: {content}"
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=temperature,
                top_p=0.95,
                top_k=40,
                max_output_tokens=8192
            )
        )
        return response.text
    except Exception as e:
        st.error(f"Gemini API error: {e}")
        return f"Error: {e}"

def call_mistral(messages: List[Dict[str, str]], temperature: float = 0.2, top_p: float = 1.0, max_tokens: int = 1024) -> str:
    payload = {
        "model": "mistral-small-latest",
        "messages": messages,
        "temperature": temperature,
        "top_p": top_p,
        "max_tokens": max_tokens,
    }

    headers = {"Authorization": f"Bearer {MISTRAL_API_KEY}", "Content-Type": "application/json"}
    resp = requests.post(MISTRAL_ENDPOINT, json=payload, headers=headers)
    resp.raise_for_status()
    data = resp.json()
    return data["choices"][0]["message"]["content"]

def call_deepseek(messages: List[Dict[str, str]], model: str = "deepseek-chat", temperature: float = 0.7, max_tokens: int = 2000) -> str:
    try:
        resp = DEEPSEEK_CLIENT.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_tokens=max_tokens,
            top_p=1.0
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"Error processing request with DeepSeek: {str(e)}"

def save_interaction(conn, feature: str, input_text: str, output_text: str):
    if st.session_state.get("username"):
        conn.execute(
            "INSERT INTO interactions (username, feature, input_text, output_text) VALUES (?, ?, ?, ?)",
            (st.session_state.username, feature, input_text, output_text),
        )
        conn.commit()

# Lease Summarization UI with OCR
def lease_summarization_ui(conn):
    st.header("📄 Lease Summary")
    if st.button("Clear Summary", key="clear_lease_summary"):
        for k in ['last_file', 'last_summary', 'last_mode', 'last_engine', 'processed_text']:
            st.session_state.pop(k, None)
        st.success("Cleared previous summary.")
        st.rerun()

    st.markdown(
        "Upload your lease PDF (scanned or text-based) and receive a concise summary. "
        "Scanned documents will be automatically processed with OCR."
    )

    uploaded_file = st.file_uploader(
        "Upload Lease Document (PDF)", type=["pdf"], key="lease_file_uploader"
    )

    if 'last_file' in st.session_state and uploaded_file:
        if st.session_state.last_file != uploaded_file.name:
            for k in ['last_summary', 'last_mode', 'last_engine', 'processed_text']:
                st.session_state.pop(k, None)

    if not uploaded_file:
        return

    # Process the PDF (with OCR if needed)
    if 'processed_text' not in st.session_state or st.session_state.get('last_file') != uploaded_file.name:
        with st.spinner("Processing document..."):
            try:
                file_bytes = uploaded_file.getvalue()
                processed_text = process_pdf_with_ocr(file_bytes, uploaded_file.name)
                st.session_state.processed_text = processed_text
                st.session_state.last_file = uploaded_file.name
            except Exception as e:
                st.error(f"Failed to process document: {str(e)}")
                return

    ai_engine = st.radio(
        "Select AI Model",
        ["In-depth"],
        index=0,
        horizontal=True,
        key="lease_ai_engine"
    )
    summary_mode = st.radio(
        "Summary Mode",
        ["Full Document", "Page-by-Page"],
        index=1,
        horizontal=True,
        key="lease_summary_mode"
    )

    # Display existing summary if available
    if 'last_summary' in st.session_state and st.session_state.get('last_file') == uploaded_file.name:
        mode = st.session_state['last_mode']
        engine = st.session_state['last_engine']
        raw = st.session_state['last_summary']

        if mode == 'Full Document':
            summary_content = raw
            st.subheader(f"Full Document Summary ({engine})")
            st.write(summary_content)
        else:
            parts = raw
            st.subheader(f"Page-by-Page Summary ({engine})")
            for idx, part in enumerate(parts, start=1):
                st.markdown(f"**Page {idx}:**")
                st.write(part)
            summary_content = "\n\n".join(parts)

        st.divider()
        st.markdown("### 📥 Export Summary")
        file_base = uploaded_file.name.rsplit(".", 1)[0]
        file_name = st.text_input("Filename (no extension):", value=file_base, key="lease_export_name")

        # Export options
        col1, col2 = st.columns(2)
        with col1:
            # PDF export
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in summary_content.split('\n'):
                pdf.multi_cell(0, 6, line)
                pdf.ln(2)
            pdf_bytes = pdf.output(dest='S').encode('latin-1')
            st.download_button(
                "Download PDF",
                data=pdf_bytes,
                file_name=f"{file_name}.pdf",
                mime="application/pdf",
                key="lease_export_pdf"
            )

        with col2:
            # Word export
            doc = docx.Document()
            doc.add_heading('Lease Summary', level=1)
            for para in summary_content.split('\n\n'):
                doc.add_paragraph(para)
            buf = io.BytesIO()
            doc.save(buf)
            st.download_button(
                "Download Word",
                data=buf.getvalue(),
                file_name=f"{file_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="lease_export_word"
            )

    # Generate new summary
    if st.button("Generate Summary", key="lease_generate_button"):
        if 'processed_text' not in st.session_state:
            st.error("Document not processed yet")
            return

        text = st.session_state.processed_text
        st.session_state['last_file'] = uploaded_file.name
        st.session_state['last_mode'] = summary_mode
        st.session_state['last_engine'] = ai_engine

        if summary_mode == "Full Document":
            with st.spinner("Summarizing full document..."):
                summaries = []
                chunks = [text[i:i+15000] for i in range(0, len(text), 15000)] if len(text) > 15000 else [text]
                for chunk in chunks:
                    prompt = (
                        "Summarize this portion of the lease agreement in clear, concise language, "
                        "preserving all key details:\n\n" + chunk
                    )
                    if ai_engine == "DeepSeek":
                        summaries.append(call_deepseek(messages=[{"role":"user","content":prompt}], model="deepseek-chat", temperature=0.3, max_tokens=1024))
                    elif ai_engine == "Gemini Pro":
                        summaries.append(call_gemini(feature="lease_analysis", content=prompt, temperature=0.3))
                    else:
                        summaries.append(call_mistral(messages=[{"role":"user","content":prompt}], temperature=0.3, max_tokens=1024))
                final = "\n\n".join(summaries)
            st.subheader("Full Document Summary")
            st.write(final)
            save_interaction(conn, "lease_summary_full", uploaded_file.name, final)
            st.session_state['last_summary'] = final
        else:
            # Split text into pages (assuming each page is separated by form feed or similar)
            pages = text.split('\f') if '\f' in text else [text]
            parts = []
            st.subheader("Page-by-Page Summaries")
            for i, pg in enumerate(pages, start=1):
                if not pg.strip():
                    parts.append("(no text detected)")
                    st.markdown(f"**Page {i}:** _(no text detected)_")
                else:
                    with st.spinner(f"Summarizing page {i}..."):
                        prompt = (
                            f"Summarize page {i} of this lease agreement in clear, concise language, covering all information:\n\n{pg}"
                        )
                        if ai_engine == "DeepSeek":
                            summary = call_deepseek(messages=[{"role":"user","content":prompt}], model="deepseek-chat", temperature=0.3, max_tokens=512)
                        elif ai_engine == "Gemini Pro":
                            summary = call_gemini(feature="lease_analysis", content=prompt, temperature=0.3)
                        else:
                            summary = call_mistral(messages=[{"role":"user","content":prompt}], temperature=0.3, max_tokens=512)
                        st.markdown(f"**Page {i} Summary:**")
                        st.write(summary)
                        parts.append(summary)
            save_interaction(conn, "lease_summary_pagewise", uploaded_file.name, json.dumps({f"page_{i}": pages[i-1] for i in range(1, len(pages)+1)}))
            st.session_state['last_summary'] = parts
        st.rerun()


def deal_structuring_ui(conn):
    """Enhanced deal structuring with persistent strategy chat until cleared."""
    st.header("💡 Creative Deal Structuring Bot")
    st.markdown("Get AI-powered strategies for your property deals")

    # Initialize session state
    if "deal_strategy_memory" not in st.session_state:
        st.session_state.deal_strategy_memory = []
        st.session_state.last_strategies = None
        st.session_state.strategy_confidences = {}  # Track confidences per strategy

    # Clear chat
    if st.button("Clear Strategies", key="clear_strategies"):
        st.session_state.deal_strategy_memory.clear()
        st.session_state.last_strategies = None
        st.session_state.strategy_confidences = {}
        st.rerun()

    # Replay chat history
    for role, msg in st.session_state.deal_strategy_memory:
        st.chat_message(role).write(msg)

    # Input form
    with st.expander("Deal Details", expanded=True):
        property_type = st.selectbox("Property Type", ["Residential", "Commercial", "Mixed-Use", "Land"])
        deal_stage = st.selectbox("Deal Stage", ["Pre-offer", "Under Contract", "Rehab Planning", "Exit Strategy"])
        financials = st.text_area("Financial Parameters")
        market_conditions = st.text_area("Market Conditions")
        special_considerations = st.text_area("Special Considerations")

    ai_model = st.radio("AI Model", ["in-depth"], horizontal=True)

    # Generate strategies
    if st.button("Generate Strategies", type="primary", key="gen_strat"):
        prompt = (
            f"Property Type: {property_type}\n"
            f"Deal Stage: {deal_stage}\n"
            f"Financial Parameters: {financials}\n"
            f"Market Conditions: {market_conditions}\n"
            f"Special Considerations: {special_considerations}\n\n"
            f"Generate strategies for this deal."
        )
        with st.spinner("Developing strategies..."):
            if ai_model == "Gemini":
                strategies = call_gemini("deal_strategy", prompt)
            elif ai_model == "Mistral":
                messages = [
                    {"role": "system", "content": "You are a real estate investment strategist. Provide creative deal structuring options."},
                    {"role": "user",   "content": prompt}
                ]
                strategies = call_mistral(messages=messages)
            else:  # DeepSeek
                messages = [
                    {"role": "system", "content": "You are an expert real estate strategist. Suggest creative deal structures with pros/cons."},
                    {"role": "user",   "content": prompt}
                ]
                strategies = call_deepseek(messages)

        # Record and display
        st.session_state.deal_strategy_memory.append(("assistant", strategies))
        st.session_state.last_strategies = strategies
        st.chat_message("assistant").write(strategies)
        st.subheader("Recommended Strategies")
        st.markdown(strategies)

        # Initialize confidences for each strategy
        matches = re.findall(
            r"Strategy\s+(\d+):\s*(.*?)(?=(?:Strategy\s+\d+:)|\Z)",
            strategies,
            flags=re.S
        )
        if matches:
            for num, _ in matches:
                strategy_key = f"Strategy {num}"
                if strategy_key not in st.session_state.strategy_confidences:
                    st.session_state.strategy_confidences[strategy_key] = 7  # Default confidence
        else:
            # Fallback if no numbered sections found
            if "Strategy 1" not in st.session_state.strategy_confidences:
                st.session_state.strategy_confidences["Strategy 1"] = 7

    # Strategy evaluation & refinement
    strategies = st.session_state.get("last_strategies")
    if strategies:
        # Parse individual strategies by number
        matches = re.findall(
            r"Strategy\s+(\d+):\s*(.*?)(?=(?:Strategy\s+\d+:)|\Z)",
            strategies,
            flags=re.S
        )
        if matches:
            strategy_dict = {f"Strategy {num}": text.strip() for num, text in matches}
        else:
            # Fallback if no numbered sections found
            strategy_dict = {"Strategy 1": strategies.strip()}

        labels = list(strategy_dict.keys())
        selected_label = st.selectbox("Which strategy do you prefer?", labels, key="eval_choice")
        selected_text = strategy_dict[selected_label]

        # Show the selected content
        st.markdown(f"**{selected_label}**")
        st.markdown(selected_text)

        # Confidence slider - gets/sets value from session state
        confidence = st.slider(
            "Confidence in this strategy",
            1, 10,
            value=st.session_state.strategy_confidences.get(selected_label, 7),
            key=f"conf_{selected_label.replace(' ', '_')}"
        )

        # Update confidence in session state
        st.session_state.strategy_confidences[selected_label] = confidence

        if st.button("Refine Strategy", key="refine_strat"):
            feedback = f"{selected_label} with confidence {confidence}/10"
            st.session_state.deal_strategy_memory.append(("user", feedback))

            refinement_prompt = (
                f"Refine this single strategy based on user feedback:\n\n"
                f"{selected_text}\n\n"
                f"Feedback: {feedback}"
            )
            if ai_model == "Gemini":
                refinement = call_gemini("deal_strategy", refinement_prompt)
            elif ai_model == "Mistral":
                messages = [
                    {"role": "system", "content": "Refine the selected strategy based on user feedback."},
                    {"role": "user",   "content": refinement_prompt}
                ]
                refinement = call_mistral(messages=messages)
            else:  # DeepSeek
                messages = [
                    {"role": "system", "content": "Refine this real estate strategy based on the provided feedback."},
                    {"role": "user",   "content": refinement_prompt}
                ]
                refinement = call_deepseek(messages)

            st.session_state.deal_strategy_memory.append(("assistant", refinement))
            st.chat_message("assistant").write(refinement)
            save_interaction(conn, "deal_strategy_refinement", selected_text, refinement)
# -------------------------------------------------offer generator----------------------------------------------------------------------------------

def build_guided_prompt(details: dict, detail_level: str) -> str:
    """
    Construct a detailed prompt from guided form data to generate a real estate purchase agreement.
    """
    buyer = details['parties']['buyer']
    buyer_rep = details['parties'].get('buyer_rep', '')
    seller = details['parties']['seller']
    seller_rep = details['parties'].get('seller_rep', '')
    buyer_line = f"- Buyer: {buyer}{f' (Represented by: {buyer_rep})' if buyer_rep else ''}"
    seller_line = f"- Seller: {seller}{f' (Represented by: {seller_rep})' if seller_rep else ''}"

    address = details['property']['address']
    county = details['property'].get('county', '')
    address_line = f"- Property Address: {address}"
    county_line = f"- County: {county}" if county else ''

    price = details['financial']['price_fmt']
    earnest = details['financial']['earnest_fmt']
    price_line = f"- Purchase Price: {price}"
    earnest_line = f"- Earnest Money Deposit: {earnest}"

    closing = details['dates']['closing']
    expiry = details['dates']['expiry']
    closing_line = f"- Proposed Closing Date: {closing}"
    expiry_line = f"- Offer Expiration: {expiry} hours from signing"

    financing = details['terms'].get('financing', '')
    contingencies = details['terms'].get('contingencies', [])
    contingencies_str = ', '.join(contingencies) if contingencies else 'None'
    special_terms = details['terms'].get('special', '')
    financing_line = f"- Financing Type: {financing}"
    contingencies_line = f"- Contingencies: {contingencies_str}"
    special_line = f"- Special Terms: {special_terms}" if special_terms else ''

    jurisdiction = details['terms'].get('jurisdiction', '')
    jurisdiction_line = f"- Governing Law: {jurisdiction}" if jurisdiction else ''

    sections = [
        "Generate a professional real estate purchase agreement with the following details:",
        buyer_line, seller_line, address_line, county_line,
        price_line, earnest_line, closing_line, expiry_line,
        financing_line, contingencies_line, special_line, jurisdiction_line,
        f"Level of Detail: {detail_level}."
    ]
    return "\n".join([s for s in sections if s])


def offer_generator_ui(conn):
    """
    Streamlit UI: Guided offer generator with buyer/seller/property details and multi-stage flow.
    PDF, Word, HTML, and Text exports format markdown-like input into styled outputs.
    """
    # Header and custom styling
    st.header("✍️ Advanced Offer Generator")
    st.markdown(
        """
        <style>
        .offer-section { background-color: #f0f2f6; border-radius: 8px; padding: 12px; margin-bottom: 16px; }
        .section-title { font-weight: bold; margin-top: 24px; }
        </style>
        """, unsafe_allow_html=True
    )

    # Initialize session state
    if 'offer_stage' not in st.session_state:
        st.session_state.update({
            'offer_stage': 'input_method',
            'offer_data': {},
            'generated_offer': None,
            'edited_offer': None,
            'review_comments': []
        })

    # Navigation bar
    stages = ["input_method", "details_entry", "offer_generation", "review_edit", "export"]
    labels = ["Input Method", "Details Entry", "Offer Generation", "Review & Edit", "Export"]
    idx = stages.index(st.session_state.offer_stage)
    cols = st.columns(len(stages))
    for i, label in enumerate(labels):
        with cols[i]:
            if i < idx:
                st.success(f"✓ {label}")
            elif i == idx:
                st.info(f"→ {label}")
            else:
                st.caption(label)

    # Stage 1: Input Method
    if st.session_state.offer_stage == 'input_method':
        st.markdown("### 1. Select Input Method")
        method = st.radio(
            "How would you like to create your offer?",
            ["Guided Form", "Free Text", "Upload Existing", "Template Library"],
            horizontal=True,
            key="offer_input_method"
        )
        st.session_state.offer_data['input_method'] = method

        with st.expander("AI Configuration"):
            ai_model = st.radio("AI Model Preference", ["in-depth"], horizontal=True, key="offer_ai_model")
            creativity = st.slider("Creativity Level", 0.0, 1.0, 0.3, key="offer_creativity")
            detail_level = st.select_slider(
                "Detail Level", ["Minimal","Standard","Comprehensive"], "Standard", key="offer_detail_level"
            )
            st.session_state.offer_data.update({
                'ai_model': ai_model,
                'creativity': creativity,
                'detail_level': detail_level
            })

        if st.button("Continue to Details", key="btn_continue_details"):
            st.session_state.offer_stage = 'details_entry'
            st.rerun()

    # Stage 2: Details Entry
    elif st.session_state.offer_stage == 'details_entry':
        st.markdown("### 2. Enter Offer Details")
        method = st.session_state.offer_data['input_method']

        if method == 'Guided Form':
            with st.form("offer_details_form"):
                # Buyer Situation
                st.markdown("**Buyer Situation**")
                st.markdown('<div class="offer-section">', unsafe_allow_html=True)
                buyer = {
                    'deposit': st.selectbox("Deposit available?", ["Low/Zero Deposit", "Enough Deposit", "Cash"], key="buyer_deposit"),
                    'credit': st.selectbox("Credit/mortgage capability", ["Excellent","Average","Poor","None"], key="buyer_credit"),
                    'experience': st.selectbox("Property investing experience", ["No Experience","Done 1-3 creative deals","Done >3 deals"], key="buyer_experience"),
                    'risk': st.selectbox("Risk appetite", ["Low – very cautious","Medium – balanced","High – very comfortable with risk"], key="buyer_risk"),
                    'goal': st.selectbox("Primary goal for property", ["Personal use","Buy-to-let","HMO","Serviced Accommodation","Social Housing","Sell ASAP"], key="buyer_goal"),
                    'timeline': st.selectbox("Preferred investment timeline", ["Few months","1-3 years",">3 years"], key="buyer_timeline")
                }
                st.markdown('</div>', unsafe_allow_html=True)

                # Seller Situation
                st.markdown("**Seller Situation**")
                st.markdown('<div class="offer-section">', unsafe_allow_html=True)
                seller = {
                    'motivation': st.text_area("Seller motivation & urgency", key="seller_motivation", height=80),
                    'financials': st.text_area("Seller financial difficulties?", key="seller_financials", height=80),
                    'arrears': st.text_input("If arrears: months behind & repossession deadline?", key="seller_arrears"),
                    'negative_equity': st.text_input("If negative equity: defer payment until market recovers?", key="seller_negative_equity")
                }
                st.markdown('</div>', unsafe_allow_html=True)

                # Property Details
                st.markdown("**Property Details**")
                st.markdown('<div class="offer-section">', unsafe_allow_html=True)
                property_info = {
                    'type_status': st.text_input("Property type & occupancy status", key="property_type"),
                    'condition': st.selectbox("Condition/refurbishment needed?", ["Great","Cosmetic","Heavy (Kitchen/Bath)","Back to Brick"], key="property_condition")
                }
                st.markdown('</div>', unsafe_allow_html=True)

                # Financial Summary
                st.markdown("**Financial Summary**")
                st.markdown('<div class="offer-section">', unsafe_allow_html=True)
                finance = {
                    'offer_price': st.text_input("Offer Price", key="offer_price_str"),
                    'market_price': st.text_input("Market Price", key="market_price_str"),
                    'cost': st.number_input("Cost", min_value=0.0, step=0.01, key="project_cost"),
                    'gdv': st.number_input("Gross Development Value", min_value=0.0, step=0.01, key="gdv"),
                    'other_details': st.text_area("Any other details?", key="other_details", height=80)
                }
                st.markdown('</div>', unsafe_allow_html=True)

                submit = st.form_submit_button("Generate Offer Draft")
                if submit:
                    missing = []
                    for field,msg in [('property_type','Property info required'),('offer_price_str','Offer Price required'),('market_price_str','Market Price required')]:
                        if not st.session_state.get(field): missing.append(msg)
                    if missing:
                        for m in missing: st.error(m)
                    else:
                        st.session_state.offer_data['details'] = {
                            'buyer': buyer,
                            'seller': seller,
                            'property': property_info,
                            'financial_summary': finance
                        }
                        st.session_state.offer_stage = 'offer_generation'
                        st.rerun()

            if st.button("← Back to Input Method", key="btn_back_to_input"):
                st.session_state.offer_stage = 'input_method'
                st.rerun()

        elif method == 'Free Text':
            text = st.text_area("Deal Details (min 50 chars)", key="offer_free_text", height=200)
            if st.button("Generate Offer Draft Free Text"):
                if len(text) < 50: st.error("Please add more detail.")
                else:
                    st.session_state.offer_data['details'] = {'free_text': text}
                    st.session_state.offer_stage = 'offer_generation'
                    st.rerun()
            if st.button("← Back to Input Method"):
                st.session_state.offer_stage = 'input_method'; st.rerun()

        elif method == 'Upload Existing':
            uploaded = st.file_uploader("Upload Document", type=["pdf","docx","txt"] )
            if uploaded and st.button("Analyze & Improve Upload"):
                text = ""
                if uploaded.type == "application/pdf":
                    reader = PdfReader(uploaded)
                    text = "\n".join(p.extract_text() or "" for p in reader.pages)
                elif uploaded.type == "text/plain":
                    text = uploaded.read().decode()
                st.session_state.offer_data['details'] = {'uploaded': text}
                st.session_state.offer_stage = 'offer_generation'; st.rerun()
            if st.button("← Back to Input Method"):
                st.session_state.offer_stage = 'input_method'; st.rerun()

        else:
            st.markdown("### Template Library")
            templates = {"Residential":"templates/standard_residential.json","Commercial":"templates/commercial_purchase.json"}
            choice = st.selectbox("Select Template", list(templates.keys()), key="offer_template")
            with st.expander("Preview"):
                try: st.json(open(templates[choice]).read())
                except: st.warning("Preview unavailable")
            if st.button("Use Template", key="btn_use_template"):
                data = json.load(open(templates[choice]))
                st.session_state.offer_data['details'] = data
                st.session_state.offer_stage = 'offer_generation'; st.rerun()
            if st.button("← Back to Input Method"):
                st.session_state.offer_stage = 'input_method'; st.rerun()

    # Stage 3: Offer Generation
    if st.session_state.offer_stage == 'offer_generation':
        d = st.session_state.offer_data
        if d['input_method'] == 'Guided Form':
            parts = ["Generate a professional purchase offer with the following details:"]
            b = d['details']['buyer']
            parts += [
                f"- Deposit: {b['deposit']}", f"- Credit: {b['credit']}", f"- Experience: {b['experience']}",
                f"- Risk Appetite: {b['risk']}", f"- Goal: {b['goal']}", f"- Timeline: {b['timeline']}"
            ]
            s = d['details']['seller']
            parts += [
                f"- Seller Motivation: {s['motivation']}", f"- Seller Financials: {s['financials']}"
            ]
            if s['arrears']: parts.append(f"- Arrears: {s['arrears']}")
            if s['negative_equity']: parts.append(f"- Negative Equity: {s['negative_equity']}")
            p = d['details']['property']
            parts.append(f"- Property: {p['type_status']} (Condition: {p['condition']})")
            fsum = d['details']['financial_summary']
            parts += [
                f"- Offer Price: {fsum['offer_price']}", f"- Market Price: {fsum['market_price']}",
                f"- Cost: {fsum['cost']}", f"- GDV: {fsum['gdv']}"
            ]
            if fsum['other_details']: parts.append(f"- Other: {fsum['other_details']}")
            prompt = "\n".join(parts)
        elif d['input_method'] == 'Free Text':
            prompt = f"Draft a purchase agreement:\n\n{d['details']['free_text']}"
        elif d['input_method'] == 'Upload Existing':
            prompt = f"Improve this draft:\n\n{d['details']['uploaded']}"
        else:
            prompt = f"Generate from template:\n\n{json.dumps(d['details'], indent=2)}"
        prompt += f"\n\nDetail Level: {d.get('detail_level','Standard')}"

        with st.spinner("Generating..."):
            offer = call_gemini('offer_generator', prompt)
            st.session_state.generated_offer = offer
            save_interaction(conn, 'offer_generator', prompt, offer)

        st.subheader("Generated Offer")
        st.markdown(offer, unsafe_allow_html=True)
        if st.button("Proceed to Review"): st.session_state.offer_stage = 'review_edit'; st.rerun()
        if st.button("← Back"): st.session_state.offer_stage = 'details_entry'; st.rerun()

    # Stage 4: Review & Edit
    if st.session_state.offer_stage == 'review_edit':
        edited = st.text_area("Edit draft", st.session_state.generated_offer, height=300, key='offer_edit')
        st.session_state.edited_offer = edited
        st.markdown("#### Comments")
        new_c = st.text_input("Add comment", key='offer_new_comment')
        if st.button("Add Comment") and new_c:
            ts = datetime.now().strftime("%Y-%m-%d %H:%M")
            st.session_state.review_comments.append({'ts':ts,'text':new_c,'resolved':False})
            st.rerun()
        for i,c in enumerate(st.session_state.review_comments):
            cols = st.columns([1,8,1])
            with cols[0]: st.markdown(f"**{c['ts']}**")
            with cols[1]: st.markdown(f"{'✓' if c['resolved'] else '◯'} {c['text']}")
            with cols[2]:
                if not c['resolved'] and st.button('Resolve', key=f'res_{i}'):
                    c['resolved']=True; st.rerun()
        if st.button('← Back'): st.session_state.offer_stage='offer_generation'; st.rerun()
        if st.button('Proceed to Export'): st.session_state.offer_stage='export'; st.rerun()

    # Stage 5: Export
    if st.session_state.offer_stage == 'export':
        content = st.session_state.edited_offer or st.session_state.generated_offer
        if st.checkbox('Include Comments', value=True):
            content += "\n\n---\n## Comments\n" + "\n".join([f"- [{c['ts']}] {c['text']}" for c in st.session_state.review_comments])
        fmt = st.selectbox('Format', ['PDF','Word','HTML','Text'], key='offer_export_format')
        name = st.text_input('File Name','purchase_offer', key='offer_export_name')

        if st.button('Download'):
            if fmt=='PDF':
                pdf=FPDF(); pdf.add_page()
                for line in content.split('\n'):
                    stripped=line.strip()
                    if stripped.startswith('**') and stripped.endswith('**'):
                        pdf.set_font('Arial','B',12); pdf.multi_cell(0,6,stripped.strip('*'))
                    elif stripped.startswith('* '):
                        pdf.set_font('Arial','',12); pdf.cell(8)
                        pdf.multi_cell(0,6,'- '+stripped.lstrip('* '))
                    else:
                        pdf.set_font('Arial','',12); pdf.multi_cell(0,6,stripped)
                st.download_button('Download PDF', pdf.output(dest='S').encode('latin-1','replace'), f"{name}.pdf", "application/pdf")
            elif fmt=='Word':
                doc=docx.Document()
                for line in content.split('\n'):
                    stripped=line.strip()
                    if stripped.startswith('**') and stripped.endswith('**'):
                        doc.add_paragraph(stripped.strip('*'), style='Heading 2')
                    elif stripped.startswith('* '): doc.add_paragraph(stripped.lstrip('* '), style='List Bullet')
                    else: doc.add_paragraph(line)
                buf=io.BytesIO(); doc.save(buf)
                st.download_button('Download Word',buf.getvalue(),f"{name}.docx",'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            elif fmt=='HTML':
                html='<html><body>'
                in_list=False
                for line in content.split('\n'):
                    if line.startswith('**') and line.endswith('**'):
                        if in_list: html+='</ul>'; in_list=False
                        html+=f"<h3>{line.strip('*')}</h3>"
                    elif line.startswith('* '):
                        if not in_list: html+='<ul>'; in_list=True
                        html+=f"<li>{line.lstrip('* ')}</li>"
                    else:
                        if in_list: html+='</ul>'; in_list=False
                        html+=f"<p>{line}</p>"
                if in_list: html+='</ul>'
                html+='</body></html>'
                st.download_button('Download HTML',html.encode('utf-8'),f"{name}.html",'text/html')
            else:
                st.download_button('Download Text',content.encode('utf-8'),f"{name}.txt",'text/plain')

        if st.button('Start New'):
            for k in list(st.session_state.keys()):
                if k.startswith('offer_'): del st.session_state[k]
            st.session_state.offer_stage='input_method'; st.rerun()
# ─── Admin Portal ──────────────────────────────────────────────────────────────
def admin_portal_ui(conn):
    """Enhanced admin portal with usage analytics and subscription management"""
    st.header("🔒 Admin Portal")

    tab1, tab2, tab3, tab4 = st.tabs(["User Management", "Subscription Management", "Content Management", "Usage Analytics"])

    with tab1:
        st.subheader("User Accounts")
        cursor = conn.cursor()
        cursor.execute("PRAGMA table_info(users)")
        columns = [column[1] for column in cursor.fetchall()]

        select_columns = ["username", "role"]
        if "last_login" in columns:
            select_columns.append("last_login")
        if "location_id" in columns:
            select_columns.append("location_id")
        if "created_at" in columns:
            select_columns.append("created_at")

        query = f"SELECT {', '.join(select_columns)} FROM users"
        users = conn.execute(query).fetchall()

        formatted_users = []
        for user in users:
            formatted_user = list(user)
            for i, col in enumerate(select_columns):
                if isinstance(formatted_user[i], str) and col in ['last_login', 'created_at']:
                    try:
                        formatted_user[i] = datetime.strptime(formatted_user[i], "%Y-%m-%d %H:%M:%S").strftime("%Y-%m-%d %H:%M")
                    except:
                        pass
            formatted_users.append(formatted_user)

        user_df = pd.DataFrame(formatted_users, columns=select_columns)
        st.dataframe(user_df)

        with st.expander("Create New User"):
            new_user = st.text_input("Username")
            new_pass = st.text_input("Password", type="password")
            user_role = st.selectbox("Role", ["user", "admin"])
            location_id = st.text_input("Location ID")

            if st.button("Add User"):
                if not new_user or not new_pass:
                    st.error("Username and password are required")
                elif len(new_pass) < 8:
                    st.error("Password must be at least 8 characters")
                else:
                    hashed = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt())
                    try:
                        if "location_id" in columns:
                            conn.execute(
                                "INSERT INTO users (username, password, role, location_id) VALUES (?, ?, ?, ?)",
                                (new_user, hashed, user_role, location_id)
                            )
                        else:
                            conn.execute(
                                "INSERT INTO users (username, password, role) VALUES (?, ?, ?)",
                                (new_user, hashed, user_role)
                            )
                        conn.commit()
                        st.success("User created successfully!")
                        time.sleep(1)
                        st.rerun()
                    except sqlite3.IntegrityError:
                        st.error("Username already exists")

    with tab2:
        st.subheader("Feature Access Control")

        users = conn.execute("SELECT username FROM users").fetchall()
        if not users:
            st.warning("No users found")
        else:
            selected_user = st.selectbox("Select User", [u[0] for u in users])

            sub = conn.execute(
                "SELECT lease_analysis, deal_structuring, offer_generator FROM subscriptions WHERE username = ?",
                (selected_user,)
            ).fetchone()

            if not sub:
                conn.execute(
                    "INSERT INTO subscriptions (username) VALUES (?)",
                    (selected_user,)
                )
                conn.commit()
                sub = (0, 0, 0)

            col1, col2, col3 = st.columns(3)
            with col1:
                lease_access = st.toggle("Lease Analysis", value=bool(sub[0]))
            with col2:
                deal_access = st.toggle("Deal Structuring", value=bool(sub[1]))
            with col3:
                offer_access = st.toggle("Offer Generator", value=bool(sub[2]))

            if st.button("Update Access"):
                conn.execute(
                    """UPDATE subscriptions
                    SET lease_analysis = ?, deal_structuring = ?, offer_generator = ?
                    WHERE username = ?""",
                    (int(lease_access), int(deal_access), int(offer_access), selected_user)
                )
                conn.commit()
                st.success("Access updated successfully!")

    with tab3:
        st.subheader("Training Content")
        with st.expander("Upload Training Materials"):
            file_type = st.selectbox("Content Type", ["Document", "Video"])
            uploaded = st.file_uploader(
                f"Upload {file_type}",
                type=["pdf", "docx", "mp4"] if file_type == "Document" else ["mp4", "mov"]
            )
            description = st.text_area("Content Description")

            if uploaded and st.button("Upload"):
                save_dir = "training_content"
                os.makedirs(save_dir, exist_ok=True)
                file_path = os.path.join(save_dir, uploaded.name)

                with open(file_path, "wb") as f:
                    f.write(uploaded.getbuffer())

                meta_path = os.path.join(save_dir, f"{uploaded.name}.meta")
                with open(meta_path, "w") as f:
                    json.dump({
                        "uploaded_by": st.session_state.username,
                        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "description": description,
                        "type": file_type.lower()
                    }, f)

                st.success(f"{file_type} uploaded successfully!")

        st.subheader("Content Library")
        if os.path.exists("training_content"):
            files = os.listdir("training_content")
            content_files = [f for f in files if not f.endswith(".meta")]

            for file in content_files:
                meta_file = f"{file}.meta"
                if meta_file in files:
                    with open(os.path.join("training_content", meta_file)) as f:
                        meta = json.load(f)
                    st.markdown(f"**{file}**")
                    st.caption(f"Type: {meta['type']} | Uploaded by: {meta['uploaded_by']}")
                    st.caption(f"Description: {meta['description']}")
                    st.download_button(
                        f"Download {file}",
                        data=open(os.path.join("training_content", file), "rb").read(),
                        file_name=file
                    )
                    st.divider()

    with tab4:
        st.subheader("Usage Analytics")
        st.write("### Feature Usage")
        usage = conn.execute(
            "SELECT feature, COUNT(*) as count FROM interactions GROUP BY feature"
        ).fetchall()
        if usage:
            fig = px.pie(
                names=[u[0] for u in usage],
                values=[u[1] for u in usage],
                title="Feature Usage Distribution"
            )
            st.plotly_chart(fig)
        else:
            st.warning("No usage data available yet")

        st.write("### User Activity")
        activity = conn.execute(
            "SELECT username, COUNT(*) as interactions "
            "FROM interactions GROUP BY username ORDER BY interactions DESC LIMIT 10"
        ).fetchall()
        if activity:
            fig = px.bar(
                x=[a[0] for a in activity],
                y=[a[1] for a in activity],
                labels={"x": "User", "y": "Interactions"},
                title="Top Users by Activity"
            )
            st.plotly_chart(fig)
        else:
            st.warning("No user activity data available")

# ─── History View ──────────────────────────────────────────────────────────────
def history_ui(conn):
    """Show user's interaction history"""
    st.header("🕒 Your History")

    if "username" not in st.session_state:
        st.warning("Please log in to view your history")
        return

    # If user has requested a full view, show it and bail out immediately
    if "current_interaction" in st.session_state:
        interaction = st.session_state.current_interaction
        st.subheader(f"Full Interaction – {interaction['timestamp']}")
        st.write(f"**Feature:** {interaction['feature']}")
        tabs = st.tabs(["Input", "Output"])
        with tabs[0]:
            st.text(interaction["input"])
        with tabs[1]:
            st.markdown(interaction["output"])

        if st.button("← Back to History"):
            del st.session_state.current_interaction
            st.rerun()
        return  # don't render the list below

    # Otherwise: render the list of past interactions
    history = conn.execute(
        "SELECT timestamp, feature, input_text, output_text "
        "FROM interactions WHERE username = ? ORDER BY timestamp DESC",
        (st.session_state.username,)
    ).fetchall()

    if not history:
        st.info("No history found – your interactions will appear here")
        return

    for i, (ts, feature, inp, out) in enumerate(history):
        with st.expander(f"{ts} • {feature}"):
            col1, col2 = st.columns(2)
            with col1:
                st.write("**Input**")
                st.text(inp[:500] + ("…" if len(inp) > 500 else ""))
            with col2:
                st.write("**Output**")
                st.text(out[:500] + ("…" if len(out) > 500 else ""))

            # single button per interaction
            if st.button(f"View Full Interaction #{i+1}", key=f"view_full_{i}"):
                st.session_state.current_interaction = {
                    "timestamp": ts,
                    "feature": feature,
                    "input": inp,
                    "output": out
                }
                st.rerun()

# ─── Chatbot Helper (Conversational) ───────────────────────────────────────────
def chatbot_ui(conn):
    """Persistent conversational chatbot beneath features"""
    if not st.session_state.get("username"):
        st.warning("Please log in to use the chatbot.")
        return
    # Retrieve conversation for this feature
    if "chat_memory" not in st.session_state:
        st.session_state["chat_memory"] = []
    st.header("🤖 AI Chatbot")
        # Clear chat button
    if st.button("Clear Chat", key="clear_chat_button"):
        st.session_state["chat_memory"] = []
    st.markdown("Chat with the assistant based on your recent output.")
    # Display past messages
    for role, message in st.session_state["chat_memory"]:
        st.chat_message(role).write(message)
    # New user message
    user_input = st.chat_input("Type your question...")
    if user_input:
        st.session_state["chat_memory"].append(("user", user_input))
        # Build context from last 10 interactions
        rows = conn.execute(
            "SELECT feature, input_text, output_text FROM interactions WHERE username=? ORDER BY timestamp DESC LIMIT 10",
            (st.session_state.username,)
        ).fetchall()
        context = "\n\n".join([f"Feature: {r[0]}\nInput: {r[1]}\nOutput: {r[2]}" for r in rows])
        prompt = f"Context:\n{context}\n\nQuestion:\n{user_input}"
        # Call AI
        if st.session_state.get("chat_model_choice", "Gemini") == "Gemini":
            answer = call_gemini("chatbot", prompt)
        elif st.session_state.get("chat_model_choice") == "Mistral":
            messages = [
                {"role": "system", "content": "You are a helpful assistant using past interactions."},
                {"role": "user", "content": prompt}
            ]
            answer = call_mistral(messages)
        else:  # DeepSeek
            messages = [
                {"role": "system", "content": "You are an AI assistant answering questions based on the user's context."},
                {"role": "user", "content": prompt}
            ]
            answer = call_deepseek(messages)
        # Append and display bot response
        st.session_state["chat_memory"].append(("assistant", answer))
        st.chat_message("assistant").write(answer)
        # Save interaction
        save_interaction(conn, "chatbot", user_input, answer)





def ocr_pdf_to_searchable(input_pdf_bytes, ocr_model=None):
    """
    Convert a non-selectable PDF (scanned document) into a searchable PDF using OCR.

    Args:
        input_pdf_bytes: Bytes of the input PDF file
        ocr_model: Optional OCR model tuple (processor, model)

    Returns:
        Bytes of the searchable PDF
    """
    from fpdf import FPDF
    from PIL import Image
    import pytesseract
    from pdf2image import convert_from_bytes

    try:
        # Convert PDF pages to images
        images = convert_from_bytes(input_pdf_bytes)

        # Create a new PDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)

        for img in images:
            # Perform OCR on each image
            if ocr_model:
                text = predict_text_with_model(img, ocr_model)
            else:
                text = pytesseract.image_to_string(img)

            # Create a new page
            pdf.add_page()

            # Add the original image
            img_path = "temp_img.jpg"
            img.save(img_path)
            pdf.image(img_path, x=10, y=8, w=190)

            # Add invisible text layer
            pdf.set_font("Arial", size=10)
            pdf.set_text_color(0, 0, 0, 0)  # Transparent text
            pdf.multi_cell(0, 5, text)

            # Clean up temp file
            os.remove(img_path)

        # Return the PDF bytes
        return pdf.output(dest='S').encode('latin-1')

    except Exception as e:
        st.error(f"OCR PDF conversion failed: {str(e)}")
        return None
def main():
    st.set_page_config(
        page_title="Property Deals AI",
        page_icon="🏠",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    st.markdown(
        f"""
        <style>
            .main {{
                background-color: {BRAND_COLORS['background']};
            }}
            .sidebar .sidebar-content {{
                background-color: {BRAND_COLORS['primary']} !important;
                color: white;
            }}
            h1, h2, h3 {{
                color: {BRAND_COLORS['primary']};
            }}
            .stButton>button {{
                background-color: {BRAND_COLORS['secondary']};
                color: white;
            }}
            .stTextInput>div>div>input,
            .stTextArea>div>div>textarea {{
                background-color: black !important;
                color: white !important;
            }}
            .stTextInput>div>div>input::placeholder,
            .stTextArea>div>div>textarea::placeholder {{
                color: white !important;
                opacity: 1 !important;
            }}
        </style>
        """,
        unsafe_allow_html=True
    )

    try:
        conn = init_db()
        create_default_admin(conn)
    except sqlite3.Error as e:
        st.error(f"Failed to initialize database: {e}")
        return

    if "logged_in" not in st.session_state:
        st.session_state.update({
            "logged_in": False,
            "username": None,
            "role": None,
            "subscription": {
                "lease_analysis": False,
                "deal_structuring": False,
                "offer_generator": False
            }
        })

    if not st.session_state.logged_in:
        login_ui(conn)
        return

    if st.session_state.logged_in:
        sub = conn.execute(
            "SELECT lease_analysis, deal_structuring, offer_generator FROM subscriptions WHERE username = ?",
            (st.session_state.username,)
        ).fetchone()

        if not sub and st.session_state.role != "admin":
            conn.execute(
                "INSERT INTO subscriptions (username) VALUES (?)",
                (st.session_state.username,)
            )
            conn.commit()
            sub = (0, 0, 0)
        elif st.session_state.role == "admin":
            sub = (1, 1, 1)

        st.session_state.subscription = {
            "lease_analysis": bool(sub[0]),
            "deal_structuring": bool(sub[1]),
            "offer_generator": bool(sub[2])
        }

    st.sidebar.title(f"Welcome, {st.session_state.username}")
    st.sidebar.markdown(f"**Location ID:** {st.session_state.get('location_id', 'Not specified')}")

    features = []

    if st.session_state.subscription.get("lease_analysis") or st.session_state.role == "admin":
        features.append("Lease Summarization")
    if st.session_state.subscription.get("deal_structuring") or st.session_state.role == "admin":
        features.append("Deal Structuring")
    if st.session_state.subscription.get("offer_generator") or st.session_state.role == "admin":
        features.append("Offer Generator")

    features.append("History")

    if st.session_state.role == "admin":
        features.insert(-1, "Admin Portal")

    selected = st.sidebar.radio("Navigation", features)

    try:
        if selected == "Lease Summarization":
            lease_summarization_ui(conn)
        elif selected == "Deal Structuring":
            deal_structuring_ui(conn)
        elif selected == "Offer Generator":
            offer_generator_ui(conn)
        elif selected == "History":
            history_ui(conn)
        elif selected == "Admin Portal" and st.session_state.role == "admin":
            admin_portal_ui(conn)
        else:
            st.error("Access Denied")
    except Exception as e:
        st.error(f"Error in {selected} feature: {e}")

    st.divider()
    chatbot_ui(conn)

    st.sidebar.divider()
    if st.sidebar.button("Logout"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

if __name__ == "__main__":
    main()